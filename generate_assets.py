import csv
import json
import math
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


FONT = ImageFont.load_default()
HEADING_FONT = None
SUBHEADING_FONT = None
BODY_FONT = None


def load_fonts():
    global HEADING_FONT, SUBHEADING_FONT, BODY_FONT
    try:
        HEADING_FONT = ImageFont.truetype("arial.ttf", 28)
        SUBHEADING_FONT = ImageFont.truetype("arial.ttf", 18)
        BODY_FONT = ImageFont.truetype("arial.ttf", 14)
    except Exception:
        HEADING_FONT = FONT
        SUBHEADING_FONT = FONT
        BODY_FONT = FONT


def draw_centered_text(draw, box, text, fill=(0, 0, 0)):
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    line_height = FONT.getbbox("Ag")[3] - FONT.getbbox("Ag")[1]
    total_height = len(lines) * line_height
    start_y = y1 + (y2 - y1 - total_height) / 2
    for i, line in enumerate(lines):
        w = draw.textlength(line, font=FONT)
        x = x1 + (x2 - x1 - w) / 2
        y = start_y + i * line_height
        draw.text((x, y), line, font=FONT, fill=fill)


def draw_box(draw, box, text, fill=(245, 245, 245), outline=(0, 0, 0)):
    draw.rectangle(box, fill=fill, outline=outline, width=2)
    draw_centered_text(draw, box, text)


def draw_dashed_rect(draw, box, dash=8, gap=6, outline=(120, 120, 120)):
    x1, y1, x2, y2 = box
    x = x1
    while x < x2:
        draw.line((x, y1, min(x + dash, x2), y1), fill=outline, width=2)
        draw.line((x, y2, min(x + dash, x2), y2), fill=outline, width=2)
        x += dash + gap
    y = y1
    while y < y2:
        draw.line((x1, y, x1, min(y + dash, y2)), fill=outline, width=2)
        draw.line((x2, y, x2, min(y + dash, y2)), fill=outline, width=2)
        y += dash + gap


def draw_arrow(draw, start, end, label=None, label_offset=(0, 0)):
    draw.line((start, end), fill=(0, 0, 0), width=2)
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = math.hypot(dx, dy) or 1
    ux, uy = dx / length, dy / length
    arrow_size = 10
    left = (end[0] - arrow_size * ux + arrow_size * uy * 0.5,
            end[1] - arrow_size * uy - arrow_size * ux * 0.5)
    right = (end[0] - arrow_size * ux - arrow_size * uy * 0.5,
             end[1] - arrow_size * uy + arrow_size * ux * 0.5)
    draw.polygon([end, left, right], fill=(0, 0, 0))

    if label:
        mid = ((start[0] + end[0]) / 2 + label_offset[0], (start[1] + end[1]) / 2 + label_offset[1])
        draw.text(mid, label, font=FONT, fill=(0, 0, 0))


# ---------------------------------------------------------------------------
# DFD Level 0 – Flipkart Context Diagram
# ---------------------------------------------------------------------------
def create_level0(path):
    w, h = 1600, 1000
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)

    # External entities
    customer = (40, 120, 240, 240)      # Customer (Mobile/Web)
    seller = (40, 420, 240, 540)        # Seller
    admin = (40, 720, 240, 840)         # Flipkart Admin

    payment = (1340, 120, 1560, 240)    # Payment Gateway (Razorpay/PayU)
    logistics = (1340, 420, 1560, 540)  # Logistics (Ekart)
    notif_ext = (1340, 720, 1560, 840)  # SMS/Email Provider

    # Central system
    system = (500, 300, 1060, 560)

    # Database
    database = (620, 700, 940, 830)

    # Trust boundary
    draw_dashed_rect(draw, (440, 240, 1120, 880))
    draw.text((450, 210), "Trust Boundary: Flipkart Internal Network", font=FONT, fill=(80, 80, 80))

    # Draw boxes
    draw_box(draw, customer, "E1 Customer\n(Mobile/Web)")
    draw_box(draw, seller, "E2 Seller")
    draw_box(draw, admin, "E3 Flipkart Admin")
    draw_box(draw, system, "Flipkart Platform\n(Microservices)")
    draw_box(draw, database, "Backend Databases\n(MySQL/Redis/ES)")
    draw_box(draw, payment, "E4 Payment\nGateway\n(Razorpay/PayU)")
    draw_box(draw, logistics, "E5 Logistics\n(Ekart)")
    draw_box(draw, notif_ext, "E6 SMS/Email\nProvider")

    # Data flows
    draw_arrow(draw, (240, 180), (500, 370), "F1 HTTPS", (0, -20))
    draw_arrow(draw, (240, 480), (500, 460), "F2 HTTPS", (0, -20))
    draw_arrow(draw, (240, 780), (500, 520), "F3 HTTPS/VPN", (0, 10))
    draw_arrow(draw, (780, 560), (780, 700), "F7 SQL/Redis", (10, 0))
    draw_arrow(draw, (1060, 380), (1340, 180), "F4 Payment API", (0, -20))
    draw_arrow(draw, (1340, 210), (1060, 420), "F4 Webhook", (0, 10))
    draw_arrow(draw, (1060, 480), (1340, 480), "F5 Shipping API", (0, -20))
    draw_arrow(draw, (1060, 530), (1340, 750), "F6 SMS/Email", (0, 10))

    img.save(path)


# ---------------------------------------------------------------------------
# DFD Level 1 – Flipkart Internal Services
# ---------------------------------------------------------------------------
def create_level1(path):
    w, h = 1800, 1200
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)

    # External entities
    customer = (30, 150, 210, 260)
    seller = (30, 500, 210, 610)
    admin = (30, 850, 210, 960)
    payment = (1560, 150, 1760, 270)
    logistics = (1560, 500, 1760, 610)
    notif_ext = (1560, 850, 1760, 960)

    # Trust boundaries
    draw_dashed_rect(draw, (260, 80, 1500, 1140))
    draw.text((270, 50), "Trust Boundary: Flipkart Microservices Network", font=FONT, fill=(80, 80, 80))

    draw_dashed_rect(draw, (1060, 130, 1480, 380), outline=(180, 60, 60))
    draw.text((1070, 110), "PCI Zone", font=FONT, fill=(180, 60, 60))

    # Processes
    p1 = (300, 140, 530, 260)   # API Gateway
    p2 = (300, 340, 530, 460)   # User & Auth Service
    p3 = (300, 540, 530, 660)   # Product & Search Service
    p4 = (300, 740, 530, 860)   # Seller Portal Service
    p5 = (680, 140, 950, 260)   # Cart & Checkout Service
    p6 = (680, 400, 950, 520)   # Order Management Service
    p7 = (1100, 170, 1440, 290) # Payment Service (PCI zone)
    p8 = (680, 660, 950, 780)   # Notification Service

    # Data stores
    d1 = (1100, 420, 1340, 510)  # D1 Users DB
    d2 = (1100, 540, 1340, 630)  # D2 Product Catalog DB
    d3 = (1100, 660, 1340, 750)  # D3 Orders DB
    d4 = (1100, 780, 1340, 870)  # D4 Payment Ledger
    d5 = (680, 870, 950, 940)    # D5 Redis Session Cache
    d6 = (1100, 900, 1340, 990)  # D6 Audit Logs

    # Draw external entities
    draw_box(draw, customer, "E1 Customer\n(Mobile/Web)")
    draw_box(draw, seller, "E2 Seller")
    draw_box(draw, admin, "E3 Admin")
    draw_box(draw, payment, "E4 Payment\nGateway")
    draw_box(draw, logistics, "E5 Logistics\n(Ekart)")
    draw_box(draw, notif_ext, "E6 SMS/Email")

    # Draw processes
    draw_box(draw, p1, "P1 API Gateway\n(Kong/Zuul)")
    draw_box(draw, p2, "P2 User & Auth\nService")
    draw_box(draw, p3, "P3 Product &\nSearch Service")
    draw_box(draw, p4, "P4 Seller Portal\nService")
    draw_box(draw, p5, "P5 Cart &\nCheckout Service")
    draw_box(draw, p6, "P6 Order Mgmt\nService")
    draw_box(draw, p7, "P7 Payment\nService")
    draw_box(draw, p8, "P8 Notification\nService")

    # Draw data stores
    draw_box(draw, d1, "D1 Users DB\n(MySQL)")
    draw_box(draw, d2, "D2 Product Catalog\n(MySQL+ES)")
    draw_box(draw, d3, "D3 Orders DB\n(MySQL)")
    draw_box(draw, d4, "D4 Payment Ledger\n(MySQL)")
    draw_box(draw, d5, "D5 Redis\nSession Cache")
    draw_box(draw, d6, "D6 Audit Logs\n(ELK)")

    # Data flows – external to API Gateway
    draw_arrow(draw, (210, 200), (300, 200), "HTTPS")
    draw_arrow(draw, (210, 555), (300, 600), "HTTPS", (0, -16))
    draw_arrow(draw, (210, 905), (300, 800), "VPN", (0, 10))

    # API Gateway to internal services
    draw_arrow(draw, (530, 220), (680, 200), "F1", (0, -16))
    draw_arrow(draw, (415, 260), (415, 340), "F2", (10, 0))
    draw_arrow(draw, (415, 460), (415, 540), "F3", (10, 0))

    # Cart/Checkout to Order Mgmt
    draw_arrow(draw, (815, 260), (815, 400), "F8", (10, 0))

    # Order Mgmt to Payment
    draw_arrow(draw, (950, 440), (1100, 230), "F9", (0, -16))

    # Payment to external gateway
    draw_arrow(draw, (1440, 230), (1560, 210), "F4", (0, -16))
    draw_arrow(draw, (1560, 240), (1440, 260), "Webhook", (0, 10))

    # Order Mgmt to Logistics
    draw_arrow(draw, (950, 480), (1560, 555), "F5", (0, -16))

    # Notification to SMS/Email
    draw_arrow(draw, (950, 740), (1560, 900), "F6", (0, -10))

    # Order Mgmt to Notification
    draw_arrow(draw, (815, 520), (815, 660), "Kafka", (10, 0))

    # Services to data stores
    draw_arrow(draw, (530, 400), (1100, 465), "SQL", (0, -16))   # Auth -> Users DB
    draw_arrow(draw, (530, 620), (1100, 580), "SQL/ES", (0, -16))  # Product -> Catalog DB
    draw_arrow(draw, (950, 460), (1100, 700), "SQL", (0, -10))   # Order -> Orders DB
    draw_arrow(draw, (1270, 290), (1270, 420), "SQL", (10, 0))   # Payment -> Payment Ledger... close enough
    draw_arrow(draw, (1220, 290), (1220, 780), "SQL", (-30, 0))  # Payment -> Payment Ledger

    # Auth -> Redis session
    draw_arrow(draw, (415, 460), (680, 905), "Session", (0, 10))

    # Services -> Audit logs
    draw_arrow(draw, (950, 500), (1100, 940), "Log", (0, 10))

    img.save(path)


# ---------------------------------------------------------------------------
# STRIDE Threat Table – Flipkart-specific
# ---------------------------------------------------------------------------
def stride_rows():
    rows = [
        {
            "ElementID": "P1",
            "Type": "Process",
            "STRIDE": "T",
            "Description": "API Gateway request smuggling",
            "AttackScenario": "Attacker crafts ambiguous HTTP requests to bypass Kong/Zuul gateway "
                              "routing rules and reach internal microservices directly, potentially "
                              "accessing admin-only endpoints.",
            "Likelihood": "Medium",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Normalize and validate HTTP requests at gateway, disable HTTP/1.0, "
                          "enforce strict parsing, and run regular smuggling tests."
        },
        {
            "ElementID": "P2",
            "Type": "Process",
            "STRIDE": "S",
            "Description": "OTP brute-force on Flipkart login",
            "AttackScenario": "Flipkart uses OTP-based mobile login. Attacker automates rapid OTP "
                              "guesses (4-6 digit) against /api/auth/verify-otp to take over "
                              "accounts without knowing the password.",
            "Likelihood": "High",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Rate-limit OTP attempts per phone number, enforce exponential backoff, "
                          "use CAPTCHA after failed attempts, and expire OTPs within 60 seconds."
        },
        {
            "ElementID": "P2",
            "Type": "Process",
            "STRIDE": "S",
            "Description": "Credential stuffing on email/password login",
            "AttackScenario": "Attacker uses leaked credential databases to automate login "
                              "attempts against Flipkart accounts that still use email/password.",
            "Likelihood": "High",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Breached-password detection, adaptive MFA, device fingerprinting, "
                          "and account lockout after repeated failures."
        },
        {
            "ElementID": "P5",
            "Type": "Process",
            "STRIDE": "T",
            "Description": "Price/coupon tampering during checkout",
            "AttackScenario": "Attacker intercepts Flipkart checkout API calls and modifies "
                              "cart total, discount amount, or coupon code in the request body "
                              "to pay less than the actual price.",
            "Likelihood": "Medium",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Server-side price recalculation from catalog DB, signed cart tokens, "
                          "coupon validation on backend, and re-verify totals before payment."
        },
        {
            "ElementID": "P6",
            "Type": "Process",
            "STRIDE": "E",
            "Description": "IDOR on order details and address endpoints",
            "AttackScenario": "Authenticated customer changes orderId or addressId in "
                              "/api/orders/{id} or /api/address/{id} to view or modify another "
                              "user's order details, delivery address, or phone number.",
            "Likelihood": "High",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Enforce object-level authorization: verify that the requesting user "
                          "owns the resource on every API call. Use UUIDs instead of sequential IDs."
        },
        {
            "ElementID": "E2",
            "Type": "External Entity",
            "STRIDE": "S",
            "Description": "Fake seller registration and listing fraud",
            "AttackScenario": "Attacker registers as a seller with stolen GST/PAN details, lists "
                              "counterfeit or non-existent products at attractive prices, collects "
                              "payments, and never ships goods.",
            "Likelihood": "Medium",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "KYC verification (video, bank account, GST validation), seller "
                          "escrow with delayed payouts, and automated fraud scoring."
        },
        {
            "ElementID": "F4",
            "Type": "Data Flow",
            "STRIDE": "S",
            "Description": "Razorpay/PayU webhook spoofing",
            "AttackScenario": "Attacker forges a payment success webhook to /api/payment/callback "
                              "causing Flipkart to mark an unpaid order as paid and ship the product.",
            "Likelihood": "Medium",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Verify webhook HMAC signatures using shared secret, validate payment "
                          "amount matches order total, and perform server-to-server verification "
                          "with Razorpay Orders API."
        },
        {
            "ElementID": "D1",
            "Type": "Data Store",
            "STRIDE": "I",
            "Description": "Mass user PII exposure (phone, Aadhaar, address)",
            "AttackScenario": "Database breach or misconfigured backup exposes 400M+ user records "
                              "including phone numbers, delivery addresses, and saved payment "
                              "methods to attackers.",
            "Likelihood": "Medium",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Encrypt sensitive columns (AES-256), use envelope encryption with "
                          "AWS KMS, restrict DB access via IAM, encrypt backups, and implement "
                          "data masking for non-production environments."
        },
        {
            "ElementID": "P3",
            "Type": "Process",
            "STRIDE": "T",
            "Description": "SQL injection on product search/filter",
            "AttackScenario": "Attacker injects SQL via search query, price range filter, or "
                              "sort parameter on /api/products/search to extract product catalog "
                              "data, seller information, or pricing strategies.",
            "Likelihood": "Medium",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Use parameterized queries for MySQL, validate Elasticsearch query DSL, "
                          "sanitize all filter inputs, and use ORM (Hibernate) consistently."
        },
        {
            "ElementID": "P5",
            "Type": "Process",
            "STRIDE": "D",
            "Description": "Flash sale DDoS (Big Billion Days)",
            "AttackScenario": "During Flipkart Big Billion Days sale, attackers or scalper bots "
                              "flood the cart/checkout APIs with millions of requests to exhaust "
                              "inventory holds and crash the checkout pipeline.",
            "Likelihood": "High",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "CDN-level DDoS protection (Akamai), rate limiting per user/IP, "
                          "queue-based checkout (virtual waiting room), bot detection via "
                          "device fingerprinting, and auto-scaling cart microservice."
        },
        {
            "ElementID": "P3",
            "Type": "Process",
            "STRIDE": "T",
            "Description": "Stored XSS in product reviews and Q&A",
            "AttackScenario": "Attacker posts a product review or Q&A answer containing "
                              "JavaScript that executes in other customers' browsers, stealing "
                              "session tokens or redirecting to phishing pages.",
            "Likelihood": "Medium",
            "Impact": "Medium",
            "Priority": "Medium",
            "Mitigation": "Server-side HTML sanitization (OWASP Java HTML Sanitizer), output "
                          "encoding, strict CSP headers, and review moderation pipeline."
        },
        {
            "ElementID": "D4",
            "Type": "Data Store",
            "STRIDE": "I",
            "Description": "Payment card data exposure (PCI DSS violation)",
            "AttackScenario": "Saved card details (CVV, full PAN) stored in plaintext in payment "
                              "ledger or application logs, leading to mass card fraud if breached.",
            "Likelihood": "Low",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Tokenize card data via Razorpay vault, never store CVV, truncate "
                          "PAN to last 4 digits, achieve PCI DSS Level 1 compliance, and "
                          "scan logs for accidental card data leakage."
        },
        {
            "ElementID": "E3",
            "Type": "External Entity",
            "STRIDE": "S",
            "Description": "Admin account phishing and session hijack",
            "AttackScenario": "Attacker sends targeted phishing email to Flipkart admin, stealing "
                              "credentials or admin session cookie to access seller payout controls, "
                              "catalog management, and user data.",
            "Likelihood": "Medium",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Hardware security keys (FIDO2) for admin login, admin-only VPN, "
                          "session binding to IP/device, privileged access management (PAM), "
                          "and just-in-time access elevation."
        },
        {
            "ElementID": "P4",
            "Type": "Process",
            "STRIDE": "T",
            "Description": "Seller inventory and price manipulation",
            "AttackScenario": "Compromised seller account or API key is used to bulk-update "
                              "product prices to 1 INR or inflate stock counts to win Buy Box "
                              "placement unfairly.",
            "Likelihood": "Medium",
            "Impact": "Medium",
            "Priority": "Medium",
            "Mitigation": "Anomaly detection on price/stock changes (>50% deviation triggers "
                          "review), seller API rate limits, and mandatory 2FA for bulk updates."
        },
        {
            "ElementID": "P6",
            "Type": "Process",
            "STRIDE": "R",
            "Description": "Delivery dispute: customer denies receipt",
            "AttackScenario": "Customer claims product was not delivered despite successful "
                              "handoff. Without signed proof of delivery, Flipkart must refund "
                              "and absorb the loss.",
            "Likelihood": "Medium",
            "Impact": "Medium",
            "Priority": "Medium",
            "Mitigation": "OTP-verified delivery, photo proof of delivery via Ekart app, "
                          "GPS-stamped delivery events, and centralized dispute audit trail."
        },
        {
            "ElementID": "D6",
            "Type": "Data Store",
            "STRIDE": "R",
            "Description": "Audit log tampering or deletion",
            "AttackScenario": "Insider or attacker with DB access deletes or modifies audit "
                              "logs to cover tracks after unauthorized data access or "
                              "fraudulent refund processing.",
            "Likelihood": "Low",
            "Impact": "Medium",
            "Priority": "Medium",
            "Mitigation": "Write-once append-only log storage, ship logs to centralized SIEM "
                          "(Splunk/ELK) in real-time, cryptographic log integrity (hash chains), "
                          "and separate log-admin role."
        },
        {
            "ElementID": "F1",
            "Type": "Data Flow",
            "STRIDE": "I",
            "Description": "Session token leakage from mobile app",
            "AttackScenario": "Flipkart Android/iOS app stores JWT or session token in "
                              "SharedPreferences/NSUserDefaults without encryption. Malware or "
                              "rooted device extracts tokens for account takeover.",
            "Likelihood": "Medium",
            "Impact": "High",
            "Priority": "High",
            "Mitigation": "Store tokens in Android Keystore / iOS Keychain, use certificate "
                          "pinning, detect rooted/jailbroken devices, and bind tokens to "
                          "device fingerprint."
        },
        {
            "ElementID": "P5",
            "Type": "Process",
            "STRIDE": "T",
            "Description": "Cart bombing / inventory hold abuse",
            "AttackScenario": "Attacker adds limited-stock flash sale items to thousands of "
                              "carts via automated scripts, holding inventory without purchasing "
                              "to deny legitimate buyers access.",
            "Likelihood": "Medium",
            "Impact": "Medium",
            "Priority": "Medium",
            "Mitigation": "Short cart reservation timeouts (5 min for flash sales), per-user "
                          "quantity limits, bot detection (CAPTCHA on add-to-cart during sales), "
                          "and inventory release on timeout."
        },
        {
            "ElementID": "P1",
            "Type": "Process",
            "STRIDE": "S",
            "Description": "Mobile app API key extraction",
            "AttackScenario": "Attacker reverse-engineers Flipkart APK to extract API keys, "
                              "client secrets, or hardcoded credentials used for internal "
                              "service authentication.",
            "Likelihood": "Medium",
            "Impact": "Medium",
            "Priority": "Medium",
            "Mitigation": "Use OAuth2 device flow instead of embedded secrets, obfuscate APK "
                          "(ProGuard/R8), rotate API keys frequently, and monitor for anomalous "
                          "API key usage patterns."
        },
        {
            "ElementID": "P7",
            "Type": "Process",
            "STRIDE": "R",
            "Description": "Refund fraud with insufficient evidence",
            "AttackScenario": "Customer or seller disputes a transaction. Without signed "
                              "receipts, delivery proof, or immutable payment logs, Flipkart "
                              "cannot prove the transaction occurred, leading to chargebacks.",
            "Likelihood": "Medium",
            "Impact": "Medium",
            "Priority": "Medium",
            "Mitigation": "Store signed payment receipts, maintain immutable transaction "
                          "ledger, integrate Razorpay dispute evidence API, and implement "
                          "automated chargeback response."
        },
    ]
    return rows


def write_csv(path, rows):
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "ElementID",
                "Type",
                "STRIDE",
                "Description",
                "AttackScenario",
                "Likelihood",
                "Impact",
                "Priority",
                "Mitigation",
            ],
        )
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def write_threat_dragon_json(path):
    data = {
        "version": "2.0.0",
        "summary": {
            "title": "Flipkart E-commerce Platform - STRIDE Threat Model",
            "owner": "Student",
            "description": "Microservices-based e-commerce platform modeled on Flipkart's "
                           "architecture: API gateway, user/auth service with OTP login, "
                           "product catalog with Elasticsearch, cart/checkout, order management, "
                           "payment service (Razorpay/PayU), seller marketplace, Ekart logistics "
                           "integration, and notification service."
        },
        "detail": {
            "diagrams": [
                {
                    "title": "DFD Level 0 - Flipkart Context Diagram",
                    "diagramType": "DFD",
                    "cells": [],
                    "size": {"width": 1600, "height": 1000}
                },
                {
                    "title": "DFD Level 1 - Flipkart Internal Services",
                    "diagramType": "DFD",
                    "cells": [],
                    "size": {"width": 1800, "height": 1200}
                }
            ],
            "threats": []
        }
    }
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


def write_tm7_placeholder(path):
    content = (
        "Placeholder file for Microsoft Threat Modeling Tool.\n"
        "System: Flipkart E-commerce Platform\n"
        "Architecture: Microservices (API Gateway, Auth, Product, Cart, Order, Payment, "
        "Seller Portal, Notification)\n"
        "Open Microsoft Threat Modeling Tool and create the model, then export to this "
        ".tm7 filename."
    )
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


def wrap_text(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = (current + " " + word).strip()
        if draw.textlength(test, font=font) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw, text, x, y, font, max_width, line_spacing):
    for line in wrap_text(draw, text, font, max_width):
        draw.text((x, y), line, font=font, fill=(0, 0, 0))
        y += line_spacing
    return y


def build_report(pdf_path, level0_path, level1_path, top_risks):
    load_fonts()
    page_w, page_h = 1240, 1754
    margin = 60
    pages = []

    def new_page():
        img = Image.new("RGB", (page_w, page_h), "white")
        return img, ImageDraw.Draw(img)

    # ---- Page 1: Title, Overview, DFD Level 0 ----
    img, draw = new_page()
    y = margin
    draw.text((margin, y), "STRIDE Threat Model - Flipkart E-commerce Platform",
              font=HEADING_FONT, fill=(0, 0, 0))
    y += 50
    draw.text((margin, y), "Assignment 1 - Secure Software Engineering",
              font=SUBHEADING_FONT, fill=(0, 0, 0))
    y += 40
    draw.text((margin, y), "System Overview", font=SUBHEADING_FONT, fill=(0, 0, 0))
    y += 28
    overview_text = (
        "This threat model analyzes Flipkart, India's leading e-commerce marketplace serving "
        "400M+ registered users and 500K+ sellers. The platform follows a microservices "
        "architecture deployed on Flipkart Cloud. Key components include: an API Gateway "
        "(Kong/Zuul) for request routing and rate limiting; a User & Auth Service supporting "
        "OTP-based mobile login and email/password authentication; a Product & Search Service "
        "backed by MySQL and Elasticsearch; a Cart & Checkout Service handling flash sales like "
        "Big Billion Days; an Order Management Service coordinating fulfillment with Ekart "
        "logistics; a Payment Service integrating Razorpay and PayU gateways within a PCI DSS "
        "compliant zone; a Seller Portal for marketplace operations; and a Notification Service "
        "for SMS, email, and push notifications."
    )
    y = draw_wrapped(draw, overview_text, margin, y, BODY_FONT, page_w - 2 * margin, 18)
    y += 10
    scope_text = (
        "In-scope: product browsing and search, user registration and OTP login, cart and "
        "checkout, order placement and tracking, seller onboarding and listing management, "
        "payment processing, and delivery coordination via Ekart. Out-of-scope: Flipkart "
        "Grocery (separate app), Myntra integration, advertising platform, and internal "
        "data analytics pipelines."
    )
    y = draw_wrapped(draw, scope_text, margin, y, BODY_FONT, page_w - 2 * margin, 18)
    y += 16
    draw.text((margin, y), "DFD Level 0 - Context Diagram", font=SUBHEADING_FONT, fill=(0, 0, 0))
    y += 26
    level0_img = Image.open(level0_path)
    level0_img.thumbnail((page_w - 2 * margin, 520))
    img.paste(level0_img, (margin, int(y)))
    pages.append(img)

    # ---- Page 2: DFD Level 1 + STRIDE Summary ----
    img, draw = new_page()
    y = margin
    draw.text((margin, y), "DFD Level 1 - Internal Microservices",
              font=SUBHEADING_FONT, fill=(0, 0, 0))
    y += 26
    level1_img = Image.open(level1_path)
    level1_img.thumbnail((page_w - 2 * margin, 600))
    img.paste(level1_img, (margin, int(y)))
    y += level1_img.size[1] + 20
    draw.text((margin, y), "STRIDE Mapping Summary", font=SUBHEADING_FONT, fill=(0, 0, 0))
    y += 26
    stride_summary = (
        "Spoofing: OTP brute-force on mobile login, credential stuffing, fake seller "
        "registration, payment webhook forgery, admin phishing, and mobile API key extraction. "
        "Tampering: price/coupon manipulation at checkout, SQL injection on search filters, "
        "stored XSS in reviews, cart bombing during flash sales, seller price manipulation, "
        "and API gateway request smuggling. "
        "Repudiation: delivery disputes without proof, refund fraud, and audit log tampering. "
        "Information Disclosure: mass PII exposure, payment card data leakage, and mobile app "
        "token extraction. "
        "Denial of Service: flash sale DDoS (Big Billion Days) and cart inventory hold abuse. "
        "Elevation of Privilege: IDOR on order/address endpoints enabling cross-user data access. "
        "Full details in STRIDE_Table.csv."
    )
    y = draw_wrapped(draw, stride_summary, margin, y, BODY_FONT, page_w - 2 * margin, 18)
    pages.append(img)

    # ---- Page 3: Scale + Top Risks 1-4 ----
    img, draw = new_page()
    y = margin
    draw.text((margin, y), "Likelihood and Impact Scale", font=SUBHEADING_FONT, fill=(0, 0, 0))
    y += 26
    scale_text = (
        "Likelihood: Low (requires insider access or advanced skill), Medium (feasible with "
        "publicly available tools and moderate effort), High (easily automated or frequently "
        "observed in the wild). "
        "Impact: Low (limited to individual user inconvenience), Medium (financial loss or "
        "data exposure affecting a subset of users), High (platform-wide outage, mass data "
        "breach, regulatory penalties under IT Act 2000 or DPDP Act 2023)."
    )
    y = draw_wrapped(draw, scale_text, margin, y, BODY_FONT, page_w - 2 * margin, 18)
    y += 16
    draw.text((margin, y), "Top 8 Prioritized Risks and Mitigations",
              font=SUBHEADING_FONT, fill=(0, 0, 0))
    y += 26

    for idx, risk in enumerate(top_risks[:4], start=1):
        title = f"{idx}. {risk['risk']}"
        draw.text((margin, y), title, font=SUBHEADING_FONT, fill=(0, 0, 0))
        y += 24
        y = draw_wrapped(draw, f"Rationale: {risk['rationale']}",
                         margin, y, BODY_FONT, page_w - 2 * margin, 18)
        y = draw_wrapped(draw, f"Short-term: {risk['short']}",
                         margin, y, BODY_FONT, page_w - 2 * margin, 18)
        y = draw_wrapped(draw, f"Long-term: {risk['long']}",
                         margin, y, BODY_FONT, page_w - 2 * margin, 18)
        y += 10
    pages.append(img)

    # ---- Page 4: Top Risks 5-8 + Assumptions ----
    img, draw = new_page()
    y = margin
    for idx, risk in enumerate(top_risks[4:], start=5):
        title = f"{idx}. {risk['risk']}"
        draw.text((margin, y), title, font=SUBHEADING_FONT, fill=(0, 0, 0))
        y += 24
        y = draw_wrapped(draw, f"Rationale: {risk['rationale']}",
                         margin, y, BODY_FONT, page_w - 2 * margin, 18)
        y = draw_wrapped(draw, f"Short-term: {risk['short']}",
                         margin, y, BODY_FONT, page_w - 2 * margin, 18)
        y = draw_wrapped(draw, f"Long-term: {risk['long']}",
                         margin, y, BODY_FONT, page_w - 2 * margin, 18)
        y += 10
    y += 6
    draw.text((margin, y), "Assumptions", font=SUBHEADING_FONT, fill=(0, 0, 0))
    y += 24
    assumptions = (
        "1. Flipkart uses a microservices architecture with an API Gateway (Kong/Zuul) as "
        "the single entry point. 2. User authentication is primarily OTP-based mobile login "
        "with email/password as secondary. 3. Payment processing is handled by external gateways "
        "(Razorpay, PayU) with Flipkart's Payment Service in a PCI DSS compliant zone. "
        "4. Ekart is the primary logistics partner with API-based integration. 5. Seller Portal "
        "is a separate service with its own authentication flow. 6. Data stores include MySQL "
        "for transactional data, Elasticsearch for product search, and Redis for session caching. "
        "7. Inter-service communication uses gRPC and Kafka event streams."
    )
    draw_wrapped(draw, assumptions, margin, y, BODY_FONT, page_w - 2 * margin, 18)
    pages.append(img)

    pages[0].save(pdf_path, "PDF", save_all=True, append_images=pages[1:])


def build_docx(docx_path, level0_path, level1_path, rows, top_risks):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_heading("STRIDE Threat Model – Flipkart E-commerce Platform", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Assignment 1 – Secure Software Engineering").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ---- System Overview ----
    doc.add_heading("1. System Overview", level=1)
    doc.add_paragraph(
        "This threat model analyzes Flipkart, India's leading e-commerce marketplace serving "
        "400M+ registered users and 500K+ sellers. The platform follows a microservices "
        "architecture deployed on Flipkart Cloud. Key components include:"
    )
    components = [
        "API Gateway (Kong/Zuul) – request routing, rate limiting, and authentication.",
        "User & Auth Service – OTP-based mobile login and email/password authentication.",
        "Product & Search Service – MySQL catalog with Elasticsearch for full-text search.",
        "Cart & Checkout Service – handles flash sales (Big Billion Days) with inventory holds.",
        "Order Management Service – coordinates fulfillment with Ekart logistics.",
        "Payment Service – integrates Razorpay and PayU gateways within a PCI DSS compliant zone.",
        "Seller Portal Service – marketplace operations, listing management, and seller KYC.",
        "Notification Service – SMS, email, and push notifications via external providers.",
    ]
    for comp in components:
        doc.add_paragraph(comp, style="List Bullet")

    doc.add_heading("Scope", level=2)
    doc.add_paragraph(
        "In-scope: product browsing and search, user registration and OTP login, cart and "
        "checkout, order placement and tracking, seller onboarding and listing management, "
        "payment processing, and delivery coordination via Ekart."
    )
    doc.add_paragraph(
        "Out-of-scope: Flipkart Grocery (separate app), Myntra integration, advertising "
        "platform, and internal data analytics pipelines."
    )

    # ---- DFD Level 0 ----
    doc.add_heading("2. DFD Level 0 – Context Diagram", level=1)
    doc.add_paragraph(
        "The context diagram shows external entities (Customer, Seller, Admin, Payment Gateway, "
        "Ekart Logistics, SMS/Email Provider) interacting with the Flipkart Platform and backend "
        "databases across a trust boundary."
    )
    doc.add_picture(level0_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- DFD Level 1 ----
    doc.add_heading("3. DFD Level 1 – Internal Microservices", level=1)
    doc.add_paragraph(
        "The Level 1 DFD decomposes the platform into eight processes (P1–P8), six data stores "
        "(D1–D6), and data flows between them. The Payment Service (P7) resides in a dedicated "
        "PCI compliance zone."
    )
    doc.add_picture(level1_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- STRIDE Mapping Summary ----
    doc.add_heading("4. STRIDE Mapping Summary", level=1)
    stride_cats = [
        ("Spoofing", "OTP brute-force on mobile login, credential stuffing, fake seller "
         "registration, payment webhook forgery, admin phishing, and mobile API key extraction."),
        ("Tampering", "Price/coupon manipulation at checkout, SQL injection on search filters, "
         "stored XSS in reviews, cart bombing during flash sales, seller price manipulation, "
         "and API gateway request smuggling."),
        ("Repudiation", "Delivery disputes without proof, refund fraud, and audit log tampering."),
        ("Information Disclosure", "Mass PII exposure, payment card data leakage (PCI DSS "
         "violation), and mobile app session token extraction."),
        ("Denial of Service", "Flash sale DDoS (Big Billion Days) and cart inventory hold abuse."),
        ("Elevation of Privilege", "IDOR on order/address endpoints enabling cross-user data access."),
    ]
    for cat, desc in stride_cats:
        p = doc.add_paragraph()
        run = p.add_run(cat + ": ")
        run.bold = True
        p.add_run(desc)

    # ---- STRIDE Threat Table ----
    doc.add_heading("5. STRIDE Threat Table", level=1)
    doc.add_paragraph("The table below lists all 20 identified threats with STRIDE classification, "
                      "attack scenarios, risk ratings, and mitigations.")

    cols = ["ElementID", "STRIDE", "Description", "AttackScenario", "Likelihood", "Impact",
            "Priority", "Mitigation"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr[i].text = col
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, col in enumerate(cols):
            row_cells[i].text = row_data[col]
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    # ---- Likelihood & Impact Scale ----
    doc.add_heading("6. Likelihood and Impact Scale", level=1)
    doc.add_paragraph(
        "Likelihood: Low (requires insider access or advanced skill), Medium (feasible with "
        "publicly available tools and moderate effort), High (easily automated or frequently "
        "observed in the wild)."
    )
    doc.add_paragraph(
        "Impact: Low (limited to individual user inconvenience), Medium (financial loss or "
        "data exposure affecting a subset of users), High (platform-wide outage, mass data "
        "breach, regulatory penalties under IT Act 2000 or DPDP Act 2023)."
    )

    # ---- Top 8 Risks ----
    doc.add_heading("7. Top 8 Prioritized Risks and Mitigations", level=1)
    for idx, risk in enumerate(top_risks, start=1):
        doc.add_heading(f"{idx}. {risk['risk']}", level=2)
        doc.add_paragraph(f"Rationale: {risk['rationale']}")
        doc.add_paragraph(f"Short-term mitigation: {risk['short']}")
        doc.add_paragraph(f"Long-term mitigation: {risk['long']}")

    # ---- Assumptions ----
    doc.add_heading("8. Assumptions", level=1)
    assumptions = [
        "Flipkart uses a microservices architecture with an API Gateway (Kong/Zuul) as the single entry point.",
        "User authentication is primarily OTP-based mobile login with email/password as secondary.",
        "Payment processing is handled by external gateways (Razorpay, PayU) with Flipkart's Payment Service in a PCI DSS compliant zone.",
        "Ekart is the primary logistics partner with API-based integration.",
        "Seller Portal is a separate service with its own authentication flow.",
        "Data stores include MySQL for transactional data, Elasticsearch for product search, and Redis for session caching.",
        "Inter-service communication uses gRPC and Kafka event streams.",
    ]
    for a in assumptions:
        doc.add_paragraph(a, style="List Number")

    doc.save(docx_path)


# ---------------------------------------------------------------------------
# Main – Generate all assignment assets
# ---------------------------------------------------------------------------
def main():
    level0_path = "DFD_Level0.png"
    level1_path = "DFD_Level1.png"
    report_path = "Report.pdf"
    csv_path = "STRIDE_Table.csv"
    td_path = "model_export.threatdragon.json"
    tm7_path = "model_export.tm7"

    create_level0(level0_path)
    create_level1(level1_path)

    rows = stride_rows()
    write_csv(csv_path, rows)
    write_threat_dragon_json(td_path)
    write_tm7_placeholder(tm7_path)

    top_risks = [
        {
            "risk": "OTP brute-force account takeover",
            "rationale": "Flipkart's OTP login (4-6 digits) is vulnerable to automated "
                         "guessing. Successful attack gives full account access including "
                         "saved addresses, payment methods, and order history for 400M+ users.",
            "short": "Rate-limit OTP attempts to 3 per phone/10 min, CAPTCHA on retry, "
                     "60-second OTP expiry.",
            "long": "Adaptive risk-based authentication, device binding, and behavioral "
                    "biometrics to detect automated attempts."
        },
        {
            "risk": "IDOR on order/address endpoints",
            "rationale": "Sequential or predictable order IDs allow any authenticated user "
                         "to access other users' orders, addresses, and phone numbers — a "
                         "direct violation of DPDP Act 2023.",
            "short": "Enforce object-level authorization on every endpoint; use UUIDs.",
            "long": "Centralized authorization policy engine (OPA), automated IDOR scanning "
                    "in CI/CD, and bug bounty program."
        },
        {
            "risk": "Price/coupon tampering at checkout",
            "rationale": "Client-side price manipulation during checkout causes direct "
                         "financial loss. During Big Billion Days with millions of "
                         "transactions, even small per-order fraud scales massively.",
            "short": "Server-side price recalculation from catalog DB, signed cart tokens.",
            "long": "Order integrity microservice with real-time fraud scoring and "
                    "automated anomaly alerting."
        },
        {
            "risk": "Razorpay/PayU webhook spoofing",
            "rationale": "Forged payment success webhooks cause Flipkart to ship products "
                         "without receiving payment. Direct revenue loss per incident.",
            "short": "Verify HMAC signatures on all webhooks, validate amounts, and "
                     "perform server-to-server payment status verification.",
            "long": "Idempotent payment reconciliation service with automated "
                    "discrepancy detection and alerting."
        },
        {
            "risk": "Flash sale DDoS (Big Billion Days)",
            "rationale": "BBD generates 10x normal traffic. Attackers or scalper bots can "
                         "amplify this to crash checkout, causing massive revenue loss and "
                         "reputational damage.",
            "short": "CDN-level DDoS protection, per-user rate limits, virtual waiting "
                     "room queue for checkout.",
            "long": "Bot detection ML model, auto-scaling microservices, and chaos "
                    "engineering for flash sale resilience testing."
        },
        {
            "risk": "Mass user PII exposure",
            "rationale": "400M+ user records (phone, address, Aadhaar) make Flipkart a "
                         "high-value target. Breach triggers DPDP Act penalties and "
                         "catastrophic trust loss.",
            "short": "Encrypt PII columns with AES-256, AWS KMS key management, and "
                     "restrict DB access via IAM policies.",
            "long": "Data classification and tokenization, DLP monitoring, and "
                    "zero-trust database access architecture."
        },
        {
            "risk": "Fake seller registration and listing fraud",
            "rationale": "Fraudulent sellers damage marketplace trust, cause financial "
                         "losses to buyers, and expose Flipkart to regulatory action.",
            "short": "Mandatory KYC with GST/PAN verification, escrow with 7-day "
                     "payout hold for new sellers.",
            "long": "ML-based seller fraud scoring, automated listing quality checks, "
                    "and cross-platform fraud intelligence sharing."
        },
        {
            "risk": "Admin account compromise",
            "rationale": "Admin access to seller payouts, catalog, and user data makes "
                         "it a high-value target. Single compromised admin can cause "
                         "platform-wide damage.",
            "short": "FIDO2 hardware keys for admin auth, VPN-only admin access, "
                     "and IP-bound sessions.",
            "long": "Privileged access management (PAM), just-in-time access elevation, "
                    "and continuous admin behavior monitoring."
        },
    ]

    build_report(report_path, level0_path, level1_path, top_risks)
    build_docx("Report.docx", level0_path, level1_path, rows, top_risks)


if __name__ == "__main__":
    main()
